import json
import os
import sys
import uuid
from pathlib import Path
from typing import List
from typing import Optional
import random
# 将项目根目录添加到 Python 路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import subprocess
import time
from datetime import datetime
import logging
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi.staticfiles import StaticFiles
from fastapi import FastAPI, HTTPException, APIRouter, Request, Form, UploadFile, File, Depends, Header
from fastapi.routing import APIRoute, Mount
from fastapi.responses import FileResponse
from fastapi.responses import JSONResponse
from openai import OpenAI
from geekaiapp.g_model import TTSRequest, VideoSubmission, AudioSubmission, VideoRequest, VideoResponse, JMRequest, \
    Item1, VideoRequest2
from geekaiapp.g_utils import save_audio_file, upload_cos, tencent_region, tencent_secret_id, \
    tencent_secret_key, tencent_bucket, siliconflow_api_url, siliconflow_auth_token, gjld_submit_video_job, \
    gjld_check_video_status, siliconflow_videomodel, siliconflow_audiomodel, siliconflow_voice, zpai_video_job, \
    zpai_check_video_status, microsoft_api_key, microsoft_base_url, ai_api_key, ai_base_url, ip, ip_tts, \
    ai_model, generate_clip, marp_path, ip_md, ip_html, port, generate_image, generate_tts, system_prompt, \
    system_prompt2, current_directory, verify_auth_token, jimeng_cookie, jimeng_sign, download_video, ip_video, \
    product_serial, ip_img, download_image
from fastapi_mcp import add_mcp_server


app = FastAPI(debug=True)

# 使用绝对路径
static_path = Path(__file__).parent / "static"
# 先挂载静态文件
app.mount("/static", StaticFiles(directory=static_path), name="static")

# 获取当前文件的绝对路径
# current_dir = os.path.dirname(os.path.abspath(__file__))
# static_dir = os.path.join(current_dir, "..", "static")
# app.mount("/static", StaticFiles(directory=static_dir), name="static")

# ooutput1 = baseurl / ip_tts
# ooutput2 = current_dir+'/static/{ip_tts}'
# print("完整路径-g_jiekou1", ooutput1)
# print("完整路径-g_jiekou2", ooutput2)

# 再挂载 APIRouter
router = APIRouter()


# 示例路由
@router.get('/hello')
def say_hello():
    return {"message": "Hello from APIRouter!"}


# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 微软服务
client = OpenAI(
    api_key=microsoft_api_key,
    base_url=microsoft_base_url
)

# 硅基流动大模型服务
client2 = OpenAI(
    api_key=ai_api_key,
    base_url=ai_base_url
)


# g微软的文生音1
@router.post('/edge/tts1')
async def generate_tts1(request_body: TTSRequest):
    try:
        data = {
            'model': request_body.model,
            'input': request_body.input,
            'voice': request_body.voice,
            'response_format': request_body.response_format,
            'speed': request_body.speed,
        }
        response = client.audio.speech.create(
            **data
        )
        filename, output_path = save_audio_file(response.content, current_directory / ip_tts)
        audio_url = f"{ip}{ip_tts}/{filename}"
        return {
            "filename": filename,
            "output_path": audio_url
        }
    except Exception as e:
        print(f"An error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# g微软的文生音，并上传到腾讯oss
@router.post("/edge/tts12")
async def generate_tts12(request_body: TTSRequest):
    """
    Generates text-to-speech audio using the edge tts API and uploads it to Tencent Cloud COS.
    """
    try:
        data = {
            'model': request_body.model,
            'input': request_body.input,
            'voice': request_body.voice,
            'response_format': request_body.response_format,
            'speed': request_body.speed,
        }
        response = client.audio.speech.create(
            **data
        )
        # Save the audio file
        filename, output_path2 = save_audio_file(response.content, current_directory / ip_tts)
        audio_url = f"{ip}{ip_tts}/{filename}"
        # Upload to COS
        etag = upload_cos('text1', tencent_region, tencent_secret_id, tencent_secret_key, tencent_bucket, filename,
                          current_directory / ip_tts)
        if etag:
            audio_url2 = f"https://{tencent_bucket}.cos.{tencent_region}.myqcloud.com/{filename}"
            return {
                "audio_url": audio_url2,
                "filename": filename,
                "output_path": audio_url,
                "etag": etag
            }
        else:
            raise HTTPException(status_code=500, detail="Failed to upload audio to COS")
    except Exception as e:
        print(f"An error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# 硅基流动-文生视频
@router.post("/gjld/video")
async def sjld_video(video_submission: VideoSubmission):
    # Use the model from the user input if provided, otherwise use the default from config
    model = video_submission.model if video_submission.model else siliconflow_videomodel

    submit_response = gjld_submit_video_job(siliconflow_api_url, siliconflow_auth_token, model, video_submission.prompt)
    if "error" in submit_response:
        raise HTTPException(status_code=500, detail=submit_response["error"])

    request_id = submit_response.get("requestId")
    if not request_id:
        raise HTTPException(status_code=500, detail="Failed to get requestId from submit response.")

    status_response = gjld_check_video_status(siliconflow_api_url, siliconflow_auth_token, request_id)
    if status_response and status_response.get("status") == "Succeed" and status_response.get("results") and \
            status_response["results"].get("videos") and len(status_response["results"]["videos"]) > 0:
        video_url = status_response['results']['videos'][0]['url']
        return {"video_url": video_url}
    else:
        raise HTTPException(status_code=500, detail="Invalid status response or missing video URL.")


# g硅基流动-文生音
@router.post("/gjld/audio")
async def gjld_audio(audio_Submission: AudioSubmission):
    audiourl = siliconflow_api_url + "/audio/speech"
    audiomodel2 = audio_Submission.model if audio_Submission.model else siliconflow_audiomodel
    voice2 = audio_Submission.voice if audio_Submission.voice else siliconflow_voice
    headers = {
        "Authorization": f"Bearer {siliconflow_auth_token}",
        "Content-Type": "application/json"
    }
    data = {
        "model": audiomodel2,
        "input": audio_Submission.input,
        "voice": voice2,
        "response_format": "mp3",
        "sample_rate": 32000,
        "stream": True,
        "speed": 1,
        "gain": 0
    }
    response = requests.post(audiourl, headers=headers, json=data)

    if response.status_code == 200:
        filename, output_path2 = save_audio_file(response.content, current_directory / ip_tts)
        audio_url = f"{ip}{ip_tts}/{filename}"
        return {
            "filename": filename,
            "output_path": audio_url
        }
        # env = 'test'
        # etag = upload_cos(env, filename, output_path)
        # return {
        #     "filename": filename,
        #     "output_path": output_path2,
        #     "etag": etag
        # }
    else:
        raise HTTPException(status_code=response.status_code, detail="请求失败")


# g智谱AI-文生视频任务的接口
@router.post("/zhipuai/video")
async def zpai_video(video_request: VideoRequest):
    """Submits a text-to-video generation job using the ZhipuAI API."""
    try:
        task_id = zpai_video_job(video_request.prompt, video_request.with_audio)
        logger.info(f"Video generation task submitted successfully. Task ID: {task_id}")

        # 检查任务状态并返回结果
        status_response = zpai_check_video_status(task_id)
        return VideoResponse(**status_response)

    except HTTPException as err:
        raise err
    except Exception as e:
        logger.error(f"An unexpected error occurred: {e}")
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {str(e)}")


# g即梦-文生图的接口 https://github.com/LLM-Red-Team/jimeng-free-api
@router.post("/jimeng/img")
async def jm_image(request1: JMRequest):
    headers = {
        "Authorization": f"Bearer {request1.image_api_key}",
        "Content-Type": "application/json",
    }
    data = {
        "model": request1.model,
        "prompt": request1.prompt,
        "negativePrompt": request1.negativePrompt,
        "width": request1.width,
        "height": request1.height,
        "sample_strength": request1.sample_strength,
    }
    try:
        response = requests.post(request1.image_generation_url, headers=headers, json=data)
        response.raise_for_status()  # 检查 HTTP 状态码
        result = response.json()
        return {
            "url": result["data"][0]["url"]
        }
    except requests.exceptions.RequestException as e:
        raise Exception(f"图像生成 API 请求失败: {e}")
    except (KeyError, IndexError) as e:
        raise Exception(f"图像生成 API 响应解析失败: {e}")


# g单词比对
@router.post("/dcbd1")
async def get_dps123(request: Request):
    try:
        data = await request.json()
        # print(f'data ={data}')
        text = data.get('content')
        # print(f'text ={text}')
        response = client2.chat.completions.create(
            model=ai_model,
            # messages=messages,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": json.dumps(text, ensure_ascii=False)
                }
            ],
            response_format={
                'type': 'json_object'
            }
        )
        # print(json.loads(response.choices[0].message.content))
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# g绘画-commfyui-bizy
# docker run -d -p 8188:8188 -v "D:/tmp/20250118/models:/app/models" -v "D:/tmp/20250118/input:/app/input" -v "D:/tmp/20250118/img:/app/temp" -v "D:/tmp/20250118/temp:/app/output/temp" -v "D:/tmp/20250118/output:/app/output" -v "D:/tmp/20250118/user:/app/user"  --name comfyui-container2 wwwzhouhui569/comfyui_bizyair:v0.4.0
@router.post("/comfyui_bizyairapi")
# async def generate_clip_endpoint(request: GenerateClipRequest,file: UploadFile = File(...)):
async def generate_clip_endpoint(prompt: str = Form(...), seed: int = Form(...), idx: int = Form(...),
                                 workflowfile: UploadFile = File(...)):
    try:
        # logger.info(f"Received request: {request}")
        logger.info(f"Received request: prompt={prompt}, seed={seed}, idx={idx}, workflowfile={workflowfile.filename}")
        # 读取上传的文件内容
        workflow_content = await workflowfile.read()
        # 将文件内容转换为json
        workflow_data = json.loads(workflow_content)
        filename, output_path, etag = generate_clip(
            prompt,
            seed,
            workflow_data,
            idx
        )
        return {
            "filename": filename,
            "output_path": output_path,
            "etag": etag
        }
    except Exception as e:
        logger.error(f"Error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# x保存上传的Markdown内容，并转换成PPT
@router.post('/pptupload')
async def upload_markdown(request: Request):
    content = await request.body()  # 异步获取请求体
    content = content.decode('utf-8')  # 将字节转换为字符串
    timestamp = str(int(time.time()))
    md_filename = f"{timestamp}.md"
    pptx_filename = f"{timestamp}.pptx"
    # 保存Markdown文件
    with open(f"{ip_md}/{md_filename}", 'w', encoding='utf-8') as f:
        f.write(content)
    # 使用marp-cli将Markdown转换为PPT
    try:
        subprocess.run([marp_path, f'{ip_md}/{md_filename}', '-o', f'{ip_html}/{pptx_filename}'], shell=True,
                       check=True)
    except subprocess.CalledProcessError as e:
        return {
            'message': 'Failed to convert Markdown to PPT',
            'error': str(e)
        }
    # 返回文件链接
    audio_url = f"{ip}{ip_md}/{pptx_filename}"
    return f'Markdown 文件已保存\n预览链接: {ip}{ip_md}/{md_filename} \n下载链接: {ip}{ip_html}/{pptx_filename}?pptx'


# g获取网页信息到word
@router.post('/generate_doc')
async def generate_doc(request: Request):
    try:
        # 获取请求中的JSON数据
        data = await request.json()
        title = data.get('title')
        content = data.get('content')

        if not title and not content:
            logger.error("Title or content is required")
            return JSONResponse({"error": "Title or content is required"}, status_code=400)

        # 生成文档
        file_name = f"llm_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(ip_html, file_name)
        logger.debug(f"File path: {file_path}")

        doc = Document()

        if title:
            # 添加大标题
            paragraph = doc.add_heading(title, level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
            paragraph.style.font.name = 'FangSong'  # 直接设置整个段落的字体
            paragraph.style.font.size = Pt(22)  # 二号字体

        if content:
            # 添加正文
            paragraph = doc.add_paragraph(content)
            paragraph.style.font.name = 'FangSong'  # 直接设置整个段落的字体
            paragraph.style.font.size = Pt(10.5)  # 五号字体

        doc.save(file_path)
        logger.info(f"Document generated successfully at {file_path}")

        # 在Mac上打开文件
        # subprocess.call(['open', file_path], shell=True)
        word_url = f"{ip}{ip_html}/{file_name}"
        return JSONResponse({"message": "Document generated successfully", "file_path": word_url}, status_code=200)

    except Exception as e:
        logger.error(f"Error generating document: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


# g把markdown2思维导图
@router.post('/markdown2map/upload')
async def upload_markdown2map(request: Request):
    content = await request.body()
    content = content.decode('utf-8')
    time_name = str(int(time.time()))  # 生成时间戳作为文件名
    md_file_name = time_name + ".md"  # Markdown文件名
    html_file_name = time_name + ".html"  # HTML文件名

    # 创建markdown和html文件夹，如果它们不存在的话
    os.makedirs('../static/markdown', exist_ok=True)
    os.makedirs('../static/html', exist_ok=True)

    # 将Markdown内容写入文件
    with open(f'{ip_md}/{md_file_name}', "w", encoding='utf-8') as f:
        f.write(content)

    print(f"Markdown file created: {ip_md}/{md_file_name}")

    # 使用subprocess调用markmap-cli将Markdown转换为HTML，并移动到static/html目录
    try:
        result = subprocess.run(['npx', 'markmap-cli', f'{ip_md}/{md_file_name}', '--no-open'], capture_output=True,
                                shell=True,
                                text=True)

        if result.returncode != 0:
            raise subprocess.CalledProcessError(result.returncode, result.args, output=result.stdout,
                                                stderr=result.stderr)

        # 尝试将生成的HTML文件移动到static/html文件夹
        os.replace(f'{ip_md}/{html_file_name}', f'{ip_html}/{html_file_name}')
        print(f"HTML file moved to: {ip_html}/{html_file_name}")

        # 返回转换后的HTML文件链接
        # return f'Markdown文件已保存. 点击预览: {request.url_for("get_html", filename=html_file_name)}'
        return f'Markdown文件已保存. 点击预览: {ip}{ip_html}/{html_file_name}'
    except subprocess.CalledProcessError as e:
        # 如果转换过程中出现错误，返回错误信息
        return f"Error generating HTML file: {e.output}\n{e.stderr}", 500


# g提供HTML文件的路径
@router.post('/static/html/{filename}')
async def get_html(filename: str):
    return FileResponse(f'../static/html/{filename}')


# g儿童绘本连读1
@router.post("/make_ai_txt_picture_audio")
async def make_ai_txt_picture_audio(data: List[Item1]):
    try:
        results = []
        for item in data:
            result = await process_data(item)
            results.append(result)
        return results
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/process-data")
async def process_data(item: Item1):
    try:
        logger.info(f"process_data 调用 generate_image API 开始")
        # 记录图像生成的开始时间
        start_time_image = time.time()
        image_url = await generate_image(item.prompt)
        # 计算图像生成耗时
        elapsed_time_image = time.time() - start_time_image
        logger.info(
            f"process_data 调用 generate_image API 结束，耗时 {elapsed_time_image:.2f} 秒，返回 image_url: {image_url}")

        logger.info(f"process_data 调用 generate_audio API 开始")
        # 记录音频生成的开始时间
        start_time_audio = time.time()
        # audio_url = await generate_audio(item.text_snippet)
        audio_url = await generate_tts(item.text_snippet, client)
        # 计算音频生成耗时
        elapsed_time_audio = time.time() - start_time_audio
        logger.info(
            f"process_data 调用 generate_audio API 结束，耗时 {elapsed_time_audio:.2f} 秒，返回 audio_url: {audio_url}")

        return {
            "description": item.text_snippet,
            "image_url": image_url,
            "audio_url": audio_url,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# json格式化输出
@router.post("/json1")
async def get_json1(request: Request):
    try:
        data = await request.json()
        # print(f'data ={data}')
        # text = data.get('content')
        # print(f'text ={text}')
        response = client2.chat.completions.create(
            model=ai_model,
            messages=[
                {
                    "role": "system",
                    "content": system_prompt2
                },
                {
                    "role": "user",
                    "content": json.dumps(data, ensure_ascii=False)
                }
            ],
            response_format={
                'type': 'json_object'
            }
        )
        print(json.loads(response.choices[0].message.content))
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# 获取当前程序运行目录并创建必要的目录
current_dir = os.path.dirname(os.path.abspath(__file__))
temp_dir = os.path.join(current_dir, 'temp')
output_dir = os.path.join(current_dir, 'static')
os.makedirs(temp_dir, exist_ok=True)
os.makedirs(output_dir, exist_ok=True)


# g把md to word
@router.post("/office/word1")
async def convert_md_to_docx(request: Request):
    from spire.doc import Document, FileFormat
    logger.info('Received request for /convert')
    content = await request.body()
    if not content:
        logger.error('No content part in the request')
        return JSONResponse(content={"error": "No content part"}, status_code=400)

    content = content.decode('utf-8')
    if content == '':
        logger.error('No content provided')
        return JSONResponse(content={"error": "No content provided"}, status_code=400)

    # 从请求的内容中读取
    mdfile_name = str(int(time.time())) + ".md"
    md_file_path = os.path.join(temp_dir, mdfile_name)
    with open(md_file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(md_file_path)

    # 创建文档实例
    doc = Document()

    # 从上传的文件加载Markdown内容
    doc.LoadFromFile(md_file_path, FileFormat.Markdown)

    # 将Markdown文件转换为Word文档并保存
    file_name = str(int(time.time())) + ".docx"
    output_path = os.path.join(output_dir, file_name)
    doc.SaveToFile(output_path, FileFormat.Docx)

    # 释放资源
    doc.Dispose()

    # 清理临时文件
    if os.path.exists(md_file_path):
        os.remove(md_file_path)

    # 返回文件的下载链接
    base_url = str(request.base_url)
    download_url = base_url + 'office/word/download/' + os.path.basename(output_path)
    print(download_url)
    # return {"download_url": download_url}
    return f'文件已保存. 点击预览: {ip}static/{file_name}'


@router.get("/office/word/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(output_dir, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path, filename=filename)


# g即梦文生视频1
@router.post("/jimeng/generate_video")
async def generate_video(request: VideoRequest2, auth_token: str = Depends(verify_auth_token)):
    try:
        logger.info(f"generate_video API 调用开始，提示词: {request.prompt}")
        start_time = time.time()

        # 从配置文件中获取视频API相关配置
        # video_api_cookie = config.get('video_api', 'cookie')
        # video_api_sign = config.get('video_api', 'sign')
        video_api_cookie = jimeng_cookie
        video_api_sign = jimeng_sign

        # 初始化视频生成API相关配置
        video_api_headers = {
            'accept': 'application/json, text/plain, */*',
            'accept-language': 'zh-CN,zh;q=0.9',
            'app-sdk-version': '48.0.0',
            'appid': '513695',
            'appvr': '5.8.0',
            'content-type': 'application/json',
            'cookie': video_api_cookie,
            'device-time': str(int(time.time())),
            'lan': 'zh-Hans',
            'loc': 'cn',
            'origin': 'https://jimeng.jianying.com',
            'pf': '7',
            'priority': 'u=1, i',
            'referer': 'https://jimeng.jianying.com/ai-tool/video/generate',
            'sec-ch-ua': '"Google Chrome";v="129", "Not=A?Brand";v="8", "Chromium";v="129"',
            'sec-ch-ua-mobile': '?0',
            'sec-ch-ua-platform': '"Windows"',
            'sec-fetch-dest': 'empty',
            'sec-fetch-mode': 'cors',
            'sec-fetch-site': 'same-origin',
            'sign': video_api_sign,
            'sign-ver': '1',
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/129.0.0.0 Safari/537.36'
        }
        video_api_base = "https://jimeng.jianying.com/mweb/v1"

        # 生成唯一的submit_id
        submit_id = str(uuid.uuid4())

        # 准备请求数据
        generate_video_payload = {
            "submit_id": submit_id,
            "task_extra": "{\"promptSource\":\"custom\",\"originSubmitId\":\"0340110f-5a94-42a9-b737-f4518f90361f\",\"isDefaultSeed\":1,\"originTemplateId\":\"\",\"imageNameMapping\":{},\"isUseAiGenPrompt\":false,\"batchNumber\":1}",
            "http_common_info": {"aid": 513695},
            "input": {
                "video_aspect_ratio": request.aspect_ratio,
                "seed": 2934141961,
                "video_gen_inputs": [
                    {
                        "prompt": request.prompt,
                        "fps": request.fps,
                        "duration_ms": request.duration_ms,
                        "video_mode": 2,
                        "template_id": ""
                    }
                ],
                "priority": 0,
                "model_req_key": "dreamina_ic_generate_video_model_vgfm_lite"
            },
            "mode": "workbench",
            "history_option": {},
            "commerce_info": {
                "resource_id": "generate_video",
                "resource_id_type": "str",
                "resource_sub_type": "aigc",
                "benefit_type": "basic_video_operation_vgfm_lite"
            },
            "client_trace_data": {}
        }

        # 发送生成视频请求
        generate_video_url = f"{video_api_base}/generate_video?aid=513695"
        logger.info(f"发送视频生成请求...")

        response = requests.post(generate_video_url, headers=video_api_headers, json=generate_video_payload)
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail=f"视频生成请求失败，状态码：{response.status_code}")

        response_data = response.json()
        if not response_data or "data" not in response_data or "aigc_data" not in response_data["data"]:
            raise HTTPException(status_code=500, detail="视频生成接口返回格式错误")

        task_id = response_data["data"]["aigc_data"]["task"]["task_id"]
        logger.info(f"视频生成任务已创建，任务ID: {task_id}")

        # 轮询检查视频生成状态
        mget_generate_task_url = f"{video_api_base}/mget_generate_task?aid=513695"
        mget_generate_task_payload = {"task_id_list": [task_id]}

        # 最多尝试30次，每次间隔2秒
        for attempt in range(30):
            time.sleep(2)
            logger.info(f"检查视频状态，第 {attempt + 1} 次尝试...")

            response = requests.post(mget_generate_task_url, headers=video_api_headers, json=mget_generate_task_payload)
            if response.status_code != 200:
                logger.warning(f"状态检查失败，状态码：{response.status_code}")
                continue

            response_data = response.json()
            if not response_data or "data" not in response_data or "task_map" not in response_data["data"]:
                logger.warning("状态检查返回格式错误")
                continue

            task_data = response_data["data"]["task_map"].get(task_id)
            if not task_data:
                logger.warning(f"未找到任务 {task_id} 的状态信息")
                continue

            task_status = task_data.get("status")
            logger.info(f"任务状态: {task_status}")

            if task_status == 50:  # 视频生成完成
                if "item_list" in task_data and task_data["item_list"] and "video" in task_data["item_list"][0]:
                    video_data = task_data["item_list"][0]["video"]
                    if "transcoded_video" in video_data and "origin" in video_data["transcoded_video"]:
                        video_url = video_data["transcoded_video"]["origin"]["video_url"]
                        elapsed_time = time.time() - start_time
                        logger.info(f"视频生成成功，耗时 {elapsed_time:.2f} 秒，URL: {video_url}")

                        # 下载视频到本地
                        try:
                            filename, file_path = download_video(video_url, current_directory / ip_video)
                            logger.info(f"视频已下载到本地: {ip}{ip_video}/{filename}")

                            # 上传到腾讯 COS
                            cos_url = upload_cos('test', tencent_region, tencent_secret_id, tencent_secret_key,
                                                 tencent_bucket, filename,
                                                 current_directory / ip_video)
                            if cos_url:
                                logger.info(f"视频已上传到 COS: {cos_url}")
                                # 删除本地文件
                                # os.remove(file_path)
                                local_url = f"{ip}{ip_video}/{filename}"
                                # return {"video_url": cos_url, "local_url": local_url, "task_id": task_id}
                                return {
                                    "video_url": cos_url,
                                    "local_url": local_url,
                                    "task_id": task_id,
                                    "markdown": f"<video controls><source src='{cos_url}' type='video/mp4'>视频预览</video>"
                                }
                            else:
                                raise HTTPException(status_code=500, detail="上传视频到 COS 失败")
                        except Exception as e:
                            logger.error(f"处理视频文件失败: {str(e)}")
                            raise HTTPException(status_code=500, detail=f"处理视频文件失败: {str(e)}")

                raise HTTPException(status_code=500, detail="视频生成完成但未找到下载地址")

        raise HTTPException(status_code=500, detail="视频生成超时")

    except Exception as e:
        logger.error(f"视频生成失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# g换脸beart1
class FaceSwapService:
    # 默认请求头
    API_HEADERS = {
        "accept": "*/*",
        "accept-language": "zh-CN,zh;q=0.9",
        "origin": "https://beart.ai",
        "priority": "u=1, i",
        "product-code": "067003",
        "referer": "https://beart.ai/",
        "sec-ch-ua": '"Google Chrome";v="129", "Not=A?Brand";v="8", "Chromium";v="129"',
        "sec-ch-ua-mobile": "?0",
        "sec-ch-ua-platform": '"Windows"',
        "sec-fetch-dest": "empty",
        "sec-fetch-mode": "cors",
        "sec-fetch-site": "same-site",
        "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/129.0.0.0 Safari/537.36"
    }

    # 从配置文件获取产品序列号
    PRODUCT_SERIAL = product_serial

    @staticmethod
    def _validate_image(image_data: bytes) -> bool:
        """验证图片格式"""
        try:
            header = image_data[:12]
            if any([
                header.startswith(b'\xFF\xD8\xFF'),  # JPEG
                header.startswith(b'\x89PNG\r\n\x1a\n'),  # PNG
                header.startswith(b'GIF87a') or header.startswith(b'GIF89a'),  # GIF
                header.startswith(b'RIFF') and header[8:12] == b'WEBP',  # WEBP
                header.startswith(b'BM')  # BMP
            ]):
                return True
            return False
        except:
            return False

    @staticmethod
    def _get_mime_type(image_data: bytes) -> str:
        """获取图片MIME类型"""
        header = image_data[:12]
        if header.startswith(b'\xFF\xD8\xFF'):
            return 'image/jpeg'
        elif header.startswith(b'\x89PNG\r\n\x1a\n'):
            return 'image/png'
        elif header.startswith(b'GIF87a') or header.startswith(b'GIF89a'):
            return 'image/gif'
        elif header.startswith(b'RIFF') and header[8:12] == b'WEBP':
            return 'image/webp'
        elif header.startswith(b'BM'):
            return 'image/bmp'
        return 'image/jpeg'

    @classmethod
    async def create_face_swap_job(cls, source_image: bytes, target_image: bytes) -> Optional[str]:
        """创建换脸任务"""
        try:
            url = "https://api.beart.ai/api/beart/face-swap/create-job"
            headers = cls.API_HEADERS.copy()
            headers.update({
                "product-serial": cls.PRODUCT_SERIAL
            })

            # 获取MIME类型
            source_mime = cls._get_mime_type(source_image)
            target_mime = cls._get_mime_type(target_image)

            # 生成随机文件名
            source_name = f"n_v{random.getrandbits(64):016x}.jpg"
            target_name = f"n_v{random.getrandbits(64):016x}.jpg"

            # 构建multipart/form-data请求
            files = {
                "target_image": (target_name, source_image, target_mime),
                "swap_image": (source_name, target_image, source_mime)
            }

            logger.info("开始上传图片...")
            response = requests.post(url, headers=headers, files=files, timeout=30)

            if response.status_code == 200:
                data = response.json()
                if data.get("code") == 100000:
                    job_id = data["result"]["job_id"]
                    logger.info(f"任务创建成功: {job_id}")
                    return job_id
                else:
                    logger.error(f"服务器返回错误: {data.get('message', {}).get('zh', '未知错误')}")
            else:
                logger.error(f"创建任务失败: HTTP {response.status_code}")
            return None

        except Exception as e:
            logger.error(f"创建任务失败: {e}")
            return None

    @classmethod
    async def get_face_swap_result(cls, job_id: str, max_retries: int = 30, interval: int = 2) -> Optional[str]:
        """获取换脸结果"""
        try:
            url = f"https://api.beart.ai/api/beart/face-swap/get-job/{job_id}"
            headers = cls.API_HEADERS.copy()
            headers["content-type"] = "application/json; charset=UTF-8"

            logger.info(f"等待处理结果，最多等待 {max_retries * interval} 秒...")
            for attempt in range(1, max_retries + 1):
                try:
                    response = requests.get(url, headers=headers, timeout=15)
                    if response.status_code == 200:
                        data = response.json()
                        if data.get("code") == 100000:
                            logger.info("处理完成")
                            return data["result"]["output"][0]
                        elif data.get("code") == 300001:  # 处理中
                            logger.info(f"处理中... {attempt}/{max_retries}")
                            time.sleep(interval)
                            continue

                    logger.error(f"获取结果失败: {response.text}")
                    return None

                except Exception as e:
                    logger.error(f"获取结果出错: {e}")
                    time.sleep(interval)

            logger.error("超过最大重试次数")
            return None

        except Exception as e:
            logger.error(f"获取结果失败: {e}")
            return None


# 修改face_swap接口，添加鉴权
@router.post("/beartAI/face-swap")
async def face_swap(
        source_image: UploadFile = File(...),
        target_image: UploadFile = File(...),
        auth_token: str = Depends(verify_auth_token)
):
    """
    换脸API接口
    - source_image: 源图片（包含要提取的人脸）
    - target_image: 目标图片（需要被替换人脸的图片）
    """
    try:
        # 读取上传的图片数据
        source_data = await source_image.read()
        target_data = await target_image.read()

        # 验证图片格式
        if not FaceSwapService._validate_image(source_data) or not FaceSwapService._validate_image(target_data):
            raise HTTPException(
                status_code=400,
                detail="不支持的图片格式，请使用jpg/png/gif/webp/bmp格式"
            )

        # 创建换脸任务
        job_id = await FaceSwapService.create_face_swap_job(source_data, target_data)
        if not job_id:
            raise HTTPException(status_code=500, detail="创建任务失败")

        # 获取处理结果
        result_url = await FaceSwapService.get_face_swap_result(job_id)
        if not result_url:
            raise HTTPException(status_code=500, detail="处理失败")

        # 下载图片并上传到COS
        try:
            # 确保输出目录存在
            # output_path = config.get('common', 'image_output_path')
            output_path = current_directory / ip_img
            if not os.path.exists(output_path):
                os.makedirs(output_path)

            # 下载图片
            filename, file_path = download_image(result_url, current_directory / ip_img)
            logger.info(f"图片已下载到本地: {file_path}")

            # 上传到腾讯 COS
            cos_url = upload_cos('test', tencent_region, tencent_secret_id, tencent_secret_key,
                                 tencent_bucket, filename,
                                 current_directory / ip_img)
            local_url = f"{ip}{ip_img}/{filename}"
            if cos_url:
                logger.info(f"图片已上传到 COS: {cos_url}")
                # 删除本地文件
                # os.remove(file_path)
                return {
                    "success": True,
                    "image_url": cos_url,
                    "original_url": local_url
                }
            else:
                raise HTTPException(status_code=500, detail="上传图片到 COS 失败")
        except Exception as e:
            logger.error(f"处理图片文件失败: {str(e)}")
            raise HTTPException(status_code=500, detail=f"处理图片文件失败: {str(e)}")

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"处理失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")




app.include_router(router, prefix="/api")

# 打印所有路由
for route in app.routes:
    if isinstance(route, APIRoute):  # 检查是否为路由
        print(f"Path: {route.path}, Methods: {route.methods}")
    elif isinstance(route, Mount):  # 检查是否为挂载点
        print(f"Mount: {route.path} -> {route.name}")

if __name__ == '__main__':
    import uvicorn

    # uvicorn.run(app, host='0.0.0.0', port=port)
    # 修改启动配置 # 禁用热重载以避免初始化问题
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info", reload=False)
